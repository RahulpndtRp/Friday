from docx import Document
from typing import List, Dict, Any, Optional
import datetime

from src.core.utils.datetime_utils import utc_now
from src.core.config.settings import Settings
from src.core.documents.models import (
    ChunkType,
    DocumentChunk,
    DocumentType,
    Path,
    ProcessedDocument,
    ProcessingStatus,
)
from src.core.documents.base_parser import BaseDocumentParser


class DOCXParser(BaseDocumentParser):
    """Parser for DOCX documents."""

    def __init__(self, settings: "Settings"):
        super().__init__(settings)
        self.supported_types = [DocumentType.DOCX]

    async def can_parse(self, file_path: str) -> bool:
        """Check if file is a DOCX."""
        extension = Path(file_path).suffix.lower()
        return extension in [".docx", ".doc"]

    async def parse_document(
        self, file_path: str, user_id: str, metadata: Optional[Dict[str, Any]] = None
    ) -> ProcessedDocument:
        """Parse DOCX document."""
        self.logger.info(f"Parsing DOCX document", file_path=file_path)

        start_time = utc_now()

        try:
            # Create document object
            doc = ProcessedDocument(
                user_id=user_id,
                original_filename=Path(file_path).name,
                file_path=file_path,
                document_type=DocumentType.DOCX,
                file_size=Path(file_path).stat().st_size,
            )

            doc.status = ProcessingStatus.PROCESSING

            # Extract metadata
            doc_metadata = await self.extract_metadata(file_path)
            doc.metadata.update(doc_metadata)

            # Extract content and create chunks
            chunks = await self._extract_docx_content(file_path, doc.id)
            doc.chunks = chunks
            doc.chunk_count = len(chunks)

            # Calculate statistics
            doc.word_count = sum(chunk.word_count for chunk in chunks)
            doc.char_count = sum(chunk.char_count for chunk in chunks)

            # Set title and classification
            doc.title = doc_metadata.get("title", Path(file_path).stem)
            doc.author = doc_metadata.get("author")
            doc.subject = doc_metadata.get("subject")
            doc.category = self._classify_document(
                doc.original_filename, " ".join(chunk.content for chunk in chunks[:3])
            )

            # Mark as completed
            doc.status = ProcessingStatus.COMPLETED
            doc.processed_at = datetime.utcnow()
            doc.processing_time = (doc.processed_at - start_time).total_seconds()

            self.logger.info(
                f"DOCX parsing completed",
                document_id=doc.id,
                chunks=len(chunks),
                processing_time=doc.processing_time,
            )

            return doc

        except Exception as e:
            self.logger.error(f"DOCX parsing failed", file_path=file_path, error=str(e))

            doc = ProcessedDocument(
                user_id=user_id,
                original_filename=Path(file_path).name,
                file_path=file_path,
                document_type=DocumentType.DOCX,
                status=ProcessingStatus.FAILED,
                error_message=str(e),
                processing_time=(utc_now() - start_time).total_seconds(),
            )
            return doc

    async def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from DOCX."""
        metadata = {}

        try:
            doc = Document(file_path)

            # Core properties
            if doc.core_properties.title:
                metadata["title"] = doc.core_properties.title
            if doc.core_properties.author:
                metadata["author"] = doc.core_properties.author
            if doc.core_properties.subject:
                metadata["subject"] = doc.core_properties.subject
            if doc.core_properties.created:
                metadata["creation_date"] = doc.core_properties.created.isoformat()
            if doc.core_properties.modified:
                metadata["modification_date"] = doc.core_properties.modified.isoformat()
            if doc.core_properties.last_modified_by:
                metadata["last_modified_by"] = doc.core_properties.last_modified_by

        except Exception as e:
            self.logger.warning(f"Failed to extract DOCX metadata", error=str(e))

        return metadata

    async def _extract_docx_content(
        self, file_path: str, document_id: str
    ) -> List[DocumentChunk]:
        """Extract content from DOCX and create chunks."""
        chunks = []

        try:
            doc = Document(file_path)

            current_section = None
            full_text = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Detect if this is a heading
                chunk_type = ChunkType.PARAGRAPH
                if para.style.name.startswith("Heading"):
                    chunk_type = ChunkType.HEADER
                    current_section = text

                full_text.append(text)

                # Create individual paragraph chunks for headers or long paragraphs
                if chunk_type == ChunkType.HEADER or len(text) > 500:
                    chunk = DocumentChunk(
                        document_id=document_id,
                        chunk_index=len(chunks),
                        chunk_type=chunk_type,
                        content=text,
                        parent_section=current_section,
                        word_count=len(text.split()),
                        char_count=len(text),
                        metadata={
                            "style": para.style.name,
                            "is_heading": chunk_type == ChunkType.HEADER,
                        },
                    )
                    chunks.append(chunk)

            # Create larger chunks from remaining text
            full_document_text = " ".join(full_text)
            text_chunks = self._chunk_text(
                full_document_text, chunk_size=800, overlap=100
            )

            for chunk_text in text_chunks:
                if not chunk_text.strip():
                    continue

                # Skip if we already have this content as individual chunks
                if any(chunk_text in chunk.content for chunk in chunks):
                    continue

                chunk = DocumentChunk(
                    document_id=document_id,
                    chunk_index=len(chunks),
                    chunk_type=ChunkType.SECTION,
                    content=chunk_text,
                    word_count=len(chunk_text.split()),
                    char_count=len(chunk_text),
                    metadata={"source": "full_document_chunk"},
                )
                chunks.append(chunk)

        except Exception as e:
            self.logger.error(f"Failed to extract DOCX content", error=str(e))
            raise

        return chunks
